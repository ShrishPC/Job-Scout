from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import os

doc = Document()

# Add Title
title = doc.add_paragraph("Job Scout: An AI-Driven Resume Parsing and Job Matching Platform")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.size = Pt(16)
title_run.font.bold = True

# Add spacing
doc.add_paragraph()

# Add Abstract heading
abstract_heading = doc.add_paragraph("Abstract")
abstract_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
abs_run = abstract_heading.runs[0]
abs_run.font.size = Pt(14)
abs_run.font.bold = True

# Add Abstract text
abstract_text = "Modern job hunting is increasingly challenging due to the high volume of applicants and the rigid reliance on automated applicant tracking systems. This project aims to develop Job Scout, an intelligent, AI-driven platform that seamlessly matches candidates to relevant job postings by analyzing their resumes using advanced Retrieval-Augmented Generation (RAG) and hybrid vector search techniques. The system integrates a Next.js frontend with a FastAPI backend, utilizing a PostgreSQL database equipped with pgvector for high-performance cosine similarity calculations. Automated web scrapers gather real-time job listings, while a local Large Language Model (Llama-3.2-3B) parses user resumes, generates high-dimensional embeddings, and tailors application assets like cover letters. The implemented platform successfully provides highly personalized job recommendations, filters opportunities based on multi-dimensional skill matrices, and visualizes candidate capabilities through an interactive radar chart, all while running efficient background updates via Celery. By automating the tedious aspects of job discovery and application tailoring, this system significantly reduces candidate fatigue and bridges the gap between employer requirements and applicant competencies in the contemporary digital recruitment landscape."
p = doc.add_paragraph(abstract_text)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Add spacing
doc.add_paragraph()

# Add Keywords
kw = doc.add_paragraph()
runner = kw.add_run("Keywords:")
runner.italic = True
kw.add_run(" Retrieval-Augmented Generation, Resume Parsing, Artificial Intelligence, Job Matching, Vector Search.")

# Save
output_dir = '/home/rishav/job-scout/weekly_modules_analysis'
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, 'Job_Scout_Abstract.docx')
doc.save(output_path)
print(f"Saved to {output_path}")

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = None # default color

def create_doc():
    doc = Document()
    
    # Title
    title = doc.add_paragraph("Job Scout: Comprehensive System Architecture & Interview Guide")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    
    doc.add_paragraph()
    
    # Section 1: Core Purpose & Market Differentiation
    add_heading(doc, "1. The Sole Purpose & Market Differentiation")
    doc.add_paragraph(
        "Why build this when tools like LinkedIn, Indeed, and existing AI resume builders exist? "
        "The sole purpose of Job Scout is absolute data sovereignty and hyper-personalization via Local AI. "
        "Unlike commercial platforms that harvest user data, sell profiles to recruiters, or charge premium "
        "fees for AI resume reviews, Job Scout runs 100% locally. No data leaves the user's machine. "
        "It is designed to give the candidate the same algorithmic power that employers use (ATS tracking, "
        "semantic search, LLMs) but reversed to serve the applicant. It acts as an autonomous agent that "
        "not only finds jobs but actively bridges the semantic gap between a resume and a job description using "
        "open-source models (Llama 3.2)."
    )
    
    # Section 2: Detailed System Diagram
    add_heading(doc, "2. System Architecture Diagram")
    diagram = """
+-------------------------------------------------------------------------+
|                              FRONTEND                                   |
|   Next.js (React) | TailwindCSS | Axios | Recharts (Admin Dashboard)    |
|   Port: 3000                                                            |
+-----------------------------------+-------------------------------------+
                                    | HTTP / REST API (CORS Allowed)
+-----------------------------------v-------------------------------------+
|                                BACKEND                                  |
|   FastAPI | Python 3.10 | Uvicorn | Port: 8001                          |
+-----------------------------------+-------------------------------------+
|   REST Controllers:                                                     |
|   - /jobs (Search, Match)                                               |
|   - /resume (Parse, Upload)                                             |
|   - /api/admin/stats (Telemetry)                                        |
+----------------+------------------+-------------------+-----------------+
                 |                  |                   |
        +--------v-------+  +-------v--------+  +-------v--------+
        |   LLM Engine   |  | Match Engine   |  | Admin Engine   |
        | (Ollama/Local) |  | (Hybrid RAG)   |  | (Telemetry)    |
        +--------+-------+  +-------+--------+  +-------+--------+
                 |                  |                   |
+----------------v------------------v-------------------v-----------------+
|                              DATABASES                                  |
|                                                                         |
|  +---------------------+                       +---------------------+  |
|  | PostgreSQL (5432)   | <---- SQLAlchemy ---- | Redis (6379)        |  |
|  | (pgvector enabled)  |                       | (Celery Broker &    |  |
|  | Tables: Jobs, Resume|                       |  FastAPI Cache)     |  |
|  +---------------------+                       +---------------------+  |
+-----------------------------------+-------------------------------------+
                                    | Celery Task Queue
+-----------------------------------v-------------------------------------+
|                            WORKER NODES                                 |
|   Celery Worker (Task execution) | Celery Beat (Cron scheduling)        |
|   Tasks: Background web scraping, Job embedding generation.             |
+-------------------------------------------------------------------------+
    """
    p = doc.add_paragraph(diagram)
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(8)
    
    # Section 3: Component Breakdown
    add_heading(doc, "3. Detailed Component Breakdown")
    
    add_heading(doc, "Frontend (Next.js & React)", level=2)
    doc.add_paragraph(
        "Built using Next.js App Router and TailwindCSS. It utilizes a neo-brutalist design language. "
        "Key features include the Candidate Dashboard (for uploading resumes and viewing matches) and the "
        "Admin Dashboard (protected by x-admin-key for viewing real-time telemetry like CPU/RAM usage and "
        "indexing velocity graphs via Recharts)."
    )
    
    add_heading(doc, "Backend API (FastAPI)", level=2)
    doc.add_paragraph(
        "A highly concurrent Python web framework. Handles file uploads (resume PDFs), sanitizes inputs, "
        "and coordinates between the database, the AI models, and the background workers. Uses dependency "
        "injection for database sessions and admin authentication."
    )
    
    add_heading(doc, "Database & Vector Store (PostgreSQL + pgvector)", level=2)
    doc.add_paragraph(
        "Acts as the primary source of truth. The pgvector extension allows the database to store "
        "high-dimensional floating-point arrays (embeddings) generated by the MiniLM model. This enables "
        "lightning-fast cosine similarity searches directly at the SQL level without needing a dedicated "
        "vector database like Pinecone."
    )
    
    add_heading(doc, "Background Tasks (Celery & Redis)", level=2)
    doc.add_paragraph(
        "Redis acts as the message broker. Celery Beat schedules periodic tasks (e.g., scraping job boards "
        "every 6 hours). Celery Workers execute these tasks asynchronously, preventing the main FastAPI "
        "thread from blocking. Redis is also used by FastAPI-Cache to cache expensive API responses."
    )
    
    add_heading(doc, "AI Engine (Llama-3.2 & MiniLM)", level=2)
    doc.add_paragraph(
        "Uses 'sentence-transformers/all-MiniLM-L6-v2' to convert text into embeddings. Uses "
        "'unsloth/Llama-3.2-3B-Instruct' (loaded via transformers) for complex natural language reasoning, "
        "such as extracting JSON skills from a raw resume or writing customized cover letters."
    )
    
    # Section 4: Tools Utilized
    add_heading(doc, "4. Tools & Technologies Utilized")
    tools = [
        ("Docker & Docker Compose", "Containerization of the entire stack ensuring environment parity. Includes 4 services: Frontend, Backend, Postgres, and Redis."),
        ("SQLAlchemy", "The Python ORM used to safely interact with PostgreSQL, utilizing parameterized queries to prevent SQL injection."),
        ("PyMuPDF (fitz)", "Used in the backend to extract raw text accurately from uploaded PDF resumes before passing it to the LLM."),
        ("Axios & Recharts", "Frontend libraries for making HTTP requests to the backend and rendering real-time telemetry graphs."),
        ("psutil", "Python library used in the backend to monitor real-time host CPU and RAM usage for the Admin Dashboard.")
    ]
    for t_name, t_desc in tools:
        p = doc.add_paragraph()
        p.add_run(f"{t_name}: ").bold = True
        p.add_run(t_desc)
        
    # Section 5: Interview FAQs
    add_heading(doc, "5. Common Interview FAQs")
    
    faqs = [
        ("Q: Why did you choose FastAPI over Django or Flask?", 
         "A: FastAPI natively supports asynchronous programming (asyncio), which is critical when dealing with I/O bound tasks like LLM inference and web scraping. It also auto-generates OpenAPI documentation and uses Pydantic for strict runtime type checking, drastically reducing bugs."),
        
        ("Q: How did you handle security in this application?", 
         "A: Although it is a local-only tool, I applied enterprise-grade security. I mitigated SQL Injection by moving entirely to SQLAlchemy parameterized queries (even in complex dynamic CASE WHEN clauses). I implemented strict MIME-type and size validation for file uploads to prevent path traversal, and I used a custom header (x-admin-key) for securing telemetry endpoints."),
         
        ("Q: Why use Postgres with pgvector instead of a dedicated Vector DB like Pinecone?",
         "A: Keeping the architecture cohesive. By using pgvector, we can perform standard relational queries (like filtering by job type or date) AND vector similarity searches in a single ACID-compliant transaction. It also eliminates the need for an external cloud dependency, keeping the app 100% local."),
         
        ("Q: How does the Hybrid RAG (Retrieval-Augmented Generation) system work?",
         "A: It uses a two-pass approach. First, it extracts hard skills from the resume and uses SQL ILIKE/Regex to do a deterministic keyword match (scoring jobs that mention exact skills). Second, it uses the MiniLM embeddings to do a semantic cosine similarity search. The scores are normalized and combined. The top results are then fed to Llama-3.2 to generate a human-readable justification for the match.")
    ]
    
    for q, a in faqs:
        p = doc.add_paragraph()
        p.add_run(q).bold = True
        doc.add_paragraph(a)
    
    output_path = '/home/rishav/job-scout/docs/Job_Scout_Comprehensive_Documentation.docx'
    import os
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Saved to {output_path}")

if __name__ == "__main__":
    create_doc()

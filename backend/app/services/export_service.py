import io
import re
import fitz
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from typing import Dict, Any

def parse_markdown_runs(text: str):
    """
    Parses a string with markdown bold formatting (**text**) into runs of (text, is_bold).
    """
    runs = []
    tokens = re.split(r'(\*\*.*?\*\*)', text)
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**') and len(token) >= 4:
            runs.append((token[2:-2], True))
        else:
            runs.append((token, False))
    return runs

def generate_docx_export(
    title: str,
    content: str,
    mode: str = "cover_letter",
    candidate_name: str | None = None,
    candidate_email: str | None = None,
    candidate_phone: str | None = None,
    company: str | None = None,
    job_title: str | None = None,
) -> bytes:
    """
    Generates a beautifully styled Word (.docx) document for a resume or cover letter.
    """
    doc = docx.Document()
    
    # Page setup: 0.75 inch margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)

    # Base style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(30, 41, 59) # Slate 800

    # Document Header
    display_name = candidate_name.strip() if candidate_name and candidate_name.strip() else "Applicant"
    p_header = doc.add_paragraph()
    p_header.paragraph_format.space_before = Pt(0)
    p_header.paragraph_format.space_after = Pt(2)
    
    r_title = p_header.add_run(display_name.upper())
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(15, 23, 42) # Slate 900

    # Contact line
    contacts = []
    if candidate_email and candidate_email.strip():
        contacts.append(candidate_email.strip())
    if candidate_phone and candidate_phone.strip():
        contacts.append(candidate_phone.strip())
    if job_title or company:
        target_str = f"Target: {job_title or 'Role'}" + (f" @ {company}" if company else "")
        contacts.append(target_str)

    p_contact = doc.add_paragraph()
    p_contact.paragraph_format.space_before = Pt(0)
    p_contact.paragraph_format.space_after = Pt(12)
    r_contact = p_contact.add_run("  •  ".join(contacts) if contacts else "Tailored Application Document")
    r_contact.font.size = Pt(9.5)
    r_contact.font.color.rgb = RGBColor(100, 116, 139) # Slate 500

    # Horizontal divider rule
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_before = Pt(0)
    p_div.paragraph_format.space_after = Pt(14)
    pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                     r'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="FA5D29"/>'
                     r'</w:pBdr>')
    p_div._p.get_or_add_pPr().append(pBdr)

    # Document Body Parsing
    lines = content.split('\n')
    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue

        # Header 1 / Section title
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(line[2:].strip().upper())
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(15, 23, 42)
            
            # Subtle section underline
            sub_bdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                r'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="CBD5E1"/>'
                                r'</w:pBdr>')
            p._p.get_or_add_pPr().append(sub_bdr)

        # Header 2 / Subsection
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(line[3:].strip())
            run.font.size = Pt(11.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(30, 41, 59)

        # Header 3 / Minor heading
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(line[4:].strip())
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(51, 65, 85)

        # Bullet item
        elif line.startswith('- ') or line.startswith('* ') or line.startswith('• '):
            bullet_text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            
            runs = parse_markdown_runs(bullet_text)
            for text_chunk, is_bold in runs:
                run = p.add_run(text_chunk)
                run.font.size = Pt(10)
                run.font.bold = is_bold
                run.font.color.rgb = RGBColor(51, 65, 85)

        # Normal Paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            
            runs = parse_markdown_runs(line)
            for text_chunk, is_bold in runs:
                run = p.add_run(text_chunk)
                run.font.size = Pt(10.5)
                run.font.bold = is_bold
                run.font.color.rgb = RGBColor(30, 41, 59)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    return doc_io.getvalue()

def generate_pdf_export(
    title: str,
    content: str,
    mode: str = "cover_letter",
    candidate_name: str | None = None,
    candidate_email: str | None = None,
    candidate_phone: str | None = None,
    company: str | None = None,
    job_title: str | None = None,
) -> bytes:
    """
    Generates a publication-grade vector PDF document using PyMuPDF.
    """
    doc = fitz.open()
    page_w, page_h = fitz.paper_size("letter")
    
    margin_l = 54
    margin_r = page_w - 54
    margin_t = 54
    margin_b = page_h - 54
    max_w = margin_r - margin_l

    font_helv = fitz.Font("helv")
    font_hebo = fitz.Font("hebo")

    def add_page(is_first: bool = False):
        page = doc.new_page(width=page_w, height=page_h)
        display_name = candidate_name.strip() if candidate_name and candidate_name.strip() else "Applicant"
        
        if is_first:
            # Header dark banner background
            page.draw_rect(fitz.Rect(0, 0, page_w, 75), color=None, fill=(0.06, 0.09, 0.16)) # Slate 950
            
            page.insert_text((margin_l, 34), display_name.upper(), fontsize=15, fontname='hebo', color=(1, 1, 1))
            
            contact_parts = []
            if candidate_email and candidate_email.strip():
                contact_parts.append(candidate_email.strip())
            if candidate_phone and candidate_phone.strip():
                contact_parts.append(candidate_phone.strip())
            if job_title or company:
                target_str = f"Target: {job_title or 'Role'}" + (f" @ {company}" if company else "")
                contact_parts.append(target_str)
            
            contact_str = "  •  ".join(contact_parts) if contact_parts else "Tailored Application Document"
            page.insert_text((margin_l, 52), contact_str, fontsize=9, fontname='helv', color=(0.75, 0.8, 0.88))
            
            # Accent bottom line
            page.draw_line(fitz.Point(0, 75), fitz.Point(page_w, 75), color=(0.98, 0.33, 0.0), width=3)
            return page, margin_t + 35
        else:
            page.insert_text((margin_l, 40), f"{display_name} — {title}", fontsize=8, fontname='helv', color=(0.5, 0.5, 0.5))
            page.draw_line(fitz.Point(margin_l, 48), fitz.Point(margin_r, 48), color=(0.85, 0.85, 0.85), width=0.75)
            return page, 65

    page, curr_y = add_page(is_first=True)

    def wrap_words(text: str, font: fitz.Font, size: float, width: float):
        words = text.split(' ')
        lines = []
        cur_line = []
        for word in words:
            test_line = ' '.join(cur_line + [word])
            if font.text_length(test_line, fontsize=size) <= width:
                cur_line.append(word)
            else:
                if cur_line:
                    lines.append(' '.join(cur_line))
                    cur_line = [word]
                else:
                    lines.append(word)
                    cur_line = []
        if cur_line:
            lines.append(' '.join(cur_line))
        return lines

    for raw_line in content.split('\n'):
        line = raw_line.strip()
        if not line:
            curr_y += 7
            continue

        if line.startswith('# '):
            curr_y += 12
            if curr_y + 25 > margin_b:
                page, curr_y = add_page(is_first=False)
            h_text = line[2:].strip()
            page.insert_text((margin_l, curr_y), h_text, fontsize=13, fontname='hebo', color=(0.06, 0.09, 0.16))
            curr_y += 4
            page.draw_line(fitz.Point(margin_l, curr_y), fitz.Point(margin_r, curr_y), color=(0.8, 0.85, 0.9), width=1)
            curr_y += 14

        elif line.startswith('## '):
            curr_y += 10
            if curr_y + 20 > margin_b:
                page, curr_y = add_page(is_first=False)
            h_text = line[3:].strip()
            page.insert_text((margin_l, curr_y), h_text, fontsize=11, fontname='hebo', color=(0.15, 0.2, 0.3))
            curr_y += 14

        elif line.startswith('### '):
            curr_y += 8
            if curr_y + 16 > margin_b:
                page, curr_y = add_page(is_first=False)
            h_text = line[4:].strip()
            page.insert_text((margin_l, curr_y), h_text, fontsize=10, fontname='hebo', color=(0.2, 0.25, 0.35))
            curr_y += 12

        elif line.startswith('- ') or line.startswith('* ') or line.startswith('• '):
            bullet_text = line[2:].strip()
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', bullet_text)
            wrapped = wrap_words(clean_text, font_helv, 9.5, max_w - 18)
            for idx, w_line in enumerate(wrapped):
                if curr_y + 14 > margin_b:
                    page, curr_y = add_page(is_first=False)
                if idx == 0:
                    page.insert_text((margin_l + 4, curr_y), '•', fontsize=10, fontname='hebo', color=(0.98, 0.33, 0.0))
                page.insert_text((margin_l + 16, curr_y), w_line, fontsize=9.5, fontname='helv', color=(0.15, 0.15, 0.15))
                curr_y += 13
            curr_y += 3

        else:
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            wrapped = wrap_words(clean_text, font_helv, 9.5, max_w)
            for w_line in wrapped:
                if curr_y + 14 > margin_b:
                    page, curr_y = add_page(is_first=False)
                page.insert_text((margin_l, curr_y), w_line, fontsize=9.5, fontname='helv', color=(0.15, 0.15, 0.15))
                curr_y += 13
            curr_y += 4

    # Page Footers
    total_pages = len(doc)
    for p_idx, p in enumerate(doc):
        p.draw_line(fitz.Point(margin_l, margin_b + 10), fitz.Point(margin_r, margin_b + 10), color=(0.85, 0.85, 0.85), width=0.5)
        p.insert_text((margin_l, margin_b + 22), 'Job Scout | AI-Powered Application Suite', fontsize=7.5, fontname='helv', color=(0.55, 0.55, 0.55))
        p.insert_text((margin_r - 55, margin_b + 22), f'Page {p_idx + 1} of {total_pages}', fontsize=7.5, fontname='helv', color=(0.55, 0.55, 0.55))

    return doc.tobytes()

def format_ats_report_markdown(ats_data: Dict[str, Any], candidate_name: str = "Candidate") -> str:
    """
    Transforms ATS evaluation JSON into a clean, human-readable Markdown report for export.
    """
    score = ats_data.get("overall_score", 0)
    rating = ats_data.get("rating_label", "EVALUATED")
    job_title = ats_data.get("job_title", "Target Position")
    company = ats_data.get("company", "Target Company")
    verdict = ats_data.get("summary_verdict", "")
    
    categories = ats_data.get("category_scores", {})
    matrix = ats_data.get("keyword_matrix", {})
    matched = matrix.get("matched_skills", [])
    missing = matrix.get("missing_skills", [])
    
    critical_missing = [s["name"] for s in missing if s.get("importance") == "critical"]
    recommended_missing = [s["name"] for s in missing if s.get("importance") != "critical"]
    matched_names = [s["name"] for s in matched]
    
    recs = ats_data.get("recommendations", [])
    suggested_rewrite = ats_data.get("suggested_bullet_rewrite", "")
    
    md_lines = [
        f"# ATS MATCH & KEYWORD GAP REPORT",
        f"**Target Role:** {job_title} @ {company}",
        f"**Candidate:** {candidate_name}",
        f"",
        f"## OVERALL SCORE: {score}% — {rating}",
        f"{verdict}",
        f"",
        f"## CATEGORY EVALUATION BREAKDOWN",
    ]
    
    for cat_key, cat_val in categories.items():
        label = cat_val.get("label", cat_key)
        c_score = cat_val.get("score", 0)
        c_max = cat_val.get("max", 0)
        md_lines.append(f"- **{label}**: {c_score} / {c_max} pts")
        
    md_lines.extend([
        f"",
        f"## KEYWORD GAP ANALYSIS",
        f"- **Matched Skills ({len(matched_names)})**: {', '.join(matched_names) if matched_names else 'None detected'}",
        f"- **Missing Critical Skills ({len(critical_missing)})**: {', '.join(critical_missing) if critical_missing else 'None — All critical skills matched!'}",
        f"- **Missing Recommended Skills ({len(recommended_missing)})**: {', '.join(recommended_missing) if recommended_missing else 'None'}",
        f"",
        f"## ACTIONABLE RECOMMENDATIONS",
    ])
    
    for r in recs:
        md_lines.append(f"- {r}")
        
    if suggested_rewrite:
        md_lines.extend([
            f"",
            f"## SUGGESTED BULLET POINT REWRITE",
            f"- {suggested_rewrite}",
        ])
        
    return "\n".join(md_lines)

def generate_ats_report_docx(ats_data: Dict[str, Any], candidate_name: str = "Applicant") -> bytes:
    """
    Generates a Word DOCX document for an ATS diagnostic report.
    """
    content = format_ats_report_markdown(ats_data, candidate_name=candidate_name)
    return generate_docx_export(
        title="ATS Compatibility Report",
        content=content,
        mode="ats_report",
        candidate_name=candidate_name,
        company=ats_data.get("company"),
        job_title=ats_data.get("job_title")
    )

def generate_ats_report_pdf(ats_data: Dict[str, Any], candidate_name: str = "Applicant") -> bytes:
    """
    Generates a vector PDF document for an ATS diagnostic report.
    """
    content = format_ats_report_markdown(ats_data, candidate_name=candidate_name)
    return generate_pdf_export(
        title="ATS Compatibility Report",
        content=content,
        mode="ats_report",
        candidate_name=candidate_name,
        company=ats_data.get("company"),
        job_title=ats_data.get("job_title")
    )
